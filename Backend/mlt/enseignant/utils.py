import io
import re
from django.http import HttpResponse

def generer_document_lecon(lecon, format_export):
    """
    Génère un HttpResponse contenant le document PDF ou Word (DOCX) d'une leçon.
    """
    format_export = format_export.lower()
    contenu_brut = lecon.contents if hasattr(lecon, 'contents') else (lecon.contenu or "Aucun contenu disponible.")
    nom_fichier = f"{lecon.titre.replace(' ', '_')}_{lecon.classe}"

    def formater_inline_pdf(texte):
        # Échapper les caractères spéciaux HTML requis par ReportLab
        texte = texte.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        # Restaurer les balises html de formatage après échappement
        # Convertir le gras **gras** -> <b>gras</b>
        texte = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', texte)
        # Convertir l'italique *italique* ou _italique_ -> <i>italique</i>
        texte = re.sub(r'\*([^*]+?)\*', r'<i>\1</i>', texte)
        texte = re.sub(r'_([^_]+?)_', r'<i>\1</i>', texte)
        # Convertir le code inline `code`
        texte = re.sub(r'`([^`]+?)`', r'<font name="Courier" color="#4F46E5"><b>\1</b></font>', texte)
        return texte

    if format_export == 'pdf':
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )

        styles = getSampleStyleSheet()
        story = []

        # Style titre principal
        titre_style = ParagraphStyle(
            'TitreLecon',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=8,
            textColor=colors.HexColor('#4F46E5'),
            fontName='Helvetica-Bold',
            alignment=TA_LEFT
        )
        
        # Style corps
        corps_style = ParagraphStyle(
            'Corps',
            parent=styles['Normal'],
            fontSize=10.5,
            leading=18,
            spaceAfter=8,
            textColor=colors.HexColor('#475569'),
            fontName='Helvetica'
        )

        # Style titres de sections (H1, H2, H3)
        h1_style = ParagraphStyle(
            'H1Style',
            parent=styles['Heading1'],
            fontSize=16,
            spaceBefore=18,
            spaceAfter=8,
            textColor=colors.HexColor('#4F46E5'),
            fontName='Helvetica-Bold'
        )
        h2_style = ParagraphStyle(
            'H2Style',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=14,
            spaceAfter=6,
            textColor=colors.HexColor('#1E293B'),
            fontName='Helvetica-Bold'
        )
        h3_style = ParagraphStyle(
            'H3Style',
            parent=styles['Heading3'],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=4,
            textColor=colors.HexColor('#475569'),
            fontName='Helvetica-Bold'
        )

        # Titre de la leçon
        story.append(Paragraph(lecon.titre, titre_style))
        story.append(Spacer(1, 4))

        # Bloc de métadonnées stylisé (Header Banner)
        meta_text = f"<b>Classe :</b> {lecon.classe}  |  <b>Thème :</b> {lecon.theme}  |  <b>Durée :</b> {lecon.duree or '45 min'}"
        meta_style = ParagraphStyle(
            'MetaStyle',
            parent=styles['Normal'],
            fontSize=9.5,
            textColor=colors.HexColor('#4F46E5'),
            fontName='Helvetica'
        )
        meta_table = Table([[Paragraph(meta_text, meta_style)]], colWidths=[17*cm])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#EEF2FF')),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#C7D2FE')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 14))

        if lecon.description:
            desc_style = ParagraphStyle(
                'DescStyle',
                parent=corps_style,
                textColor=colors.HexColor('#6B7280'),
                fontName='Helvetica-Oblique'
            )
            story.append(Paragraph(lecon.description, desc_style))
            story.append(Spacer(1, 8))

        # Parser le contenu Markdown
        in_blockquote = False
        blockquote_lines = []

        for ligne in contenu_brut.split('\n'):
            ligne_stripped = ligne.strip()
            
            if in_blockquote and not ligne_stripped.startswith('>'):
                # Générer le callout blockquote accumulé
                citation_texte = "<br/>".join(blockquote_lines)
                blockquote_style = ParagraphStyle(
                    'BlockQuoteStyle',
                    parent=styles['Normal'],
                    fontSize=9.5,
                    textColor=colors.HexColor('#5B21B6'),
                    fontName='Helvetica-Oblique',
                    leading=16
                )
                quote_table = Table([[Paragraph(citation_texte, blockquote_style)]], colWidths=[17*cm])
                quote_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F3FF')),
                    ('LINELEFT', (0, 0), (0, -1), 4, colors.HexColor('#4F46E5')),
                    ('TOPPADDING', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ('LEFTPADDING', (0, 0), (-1, -1), 14),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 14),
                ]))
                story.append(Spacer(1, 8))
                story.append(quote_table)
                story.append(Spacer(1, 8))
                blockquote_lines = []
                in_blockquote = False

            if not ligne_stripped:
                if not in_blockquote:
                    story.append(Spacer(1, 6))
                else:
                    blockquote_lines.append("")
            elif ligne_stripped.startswith('>'):
                in_blockquote = True
                texte_quote = formater_inline_pdf(ligne_stripped.lstrip('>').strip())
                blockquote_lines.append(texte_quote)
            elif ligne_stripped.startswith('### '):
                story.append(Paragraph(formater_inline_pdf(ligne_stripped[4:]), h3_style))
            elif ligne_stripped.startswith('## '):
                story.append(Paragraph(formater_inline_pdf(ligne_stripped[3:]), h2_style))
            elif ligne_stripped.startswith('# '):
                story.append(Paragraph(formater_inline_pdf(ligne_stripped[2:]), h1_style))
            elif ligne_stripped.startswith('- ') or ligne_stripped.startswith('* '):
                story.append(Paragraph(f"• {formater_inline_pdf(ligne_stripped[2:])}", corps_style))
            else:
                story.append(Paragraph(formater_inline_pdf(ligne_stripped), corps_style))

        # Si le fichier finit par une citation
        if in_blockquote and blockquote_lines:
            citation_texte = "<br/>".join(blockquote_lines)
            blockquote_style = ParagraphStyle(
                'BlockQuoteStyleEnd',
                parent=styles['Normal'],
                fontSize=9.5,
                textColor=colors.HexColor('#5B21B6'),
                fontName='Helvetica-Oblique',
                leading=16
            )
            quote_table = Table([[Paragraph(citation_texte, blockquote_style)]], colWidths=[17*cm])
            quote_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F5F3FF')),
                ('LINELEFT', (0, 0), (0, -1), 4, colors.HexColor('#4F46E5')),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('LEFTPADDING', (0, 0), (-1, -1), 14),
                ('RIGHTPADDING', (0, 0), (-1, -1), 14),
            ]))
            story.append(Spacer(1, 8))
            story.append(quote_table)

        doc.build(story)
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.pdf"'
        return response

    elif format_export == 'docx':
        import docx
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc_word = docx.Document()

        # Marges standardisées
        for sec in doc_word.sections:
            sec.top_margin = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)

        # Titre de la leçon
        p_titre = doc_word.add_paragraph()
        p_titre.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_titre.paragraph_format.space_after = Pt(4)
        run_titre = p_titre.add_run(lecon.titre)
        run_titre.font.name = 'Arial'
        run_titre.font.size = Pt(22)
        run_titre.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        run_titre.bold = True

        # Bloc Métadonnées sous forme de Callout Box (Tableau mono-cellule)
        meta_table = doc_word.add_table(rows=1, cols=1)
        meta_table.autofit = False
        meta_table.columns[0].width = Cm(16)
        meta_cell = meta_table.cell(0, 0)
        
        # Shading & border XML
        meta_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEF2FF"/>')
        meta_cell._tc.get_or_add_tcPr().append(meta_shading)
        meta_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/><w:top w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/><w:right w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="C7D2FE"/></w:tcBorders>')
        meta_cell._tc.get_or_add_tcPr().append(meta_borders)

        p_meta = meta_cell.paragraphs[0]
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_meta = p_meta.add_run(f"Niveau : {lecon.classe}   |   Thème : {lecon.theme}   |   Durée : {lecon.duree or '45 min'}")
        run_meta.font.name = 'Arial'
        run_meta.font.size = Pt(9.5)
        run_meta.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        run_meta.bold = True

        doc_word.add_paragraph() # Espacement après métadonnées

        if lecon.description:
            p_desc = doc_word.add_paragraph()
            p_desc.paragraph_format.space_after = Pt(12)
            run_desc = p_desc.add_run(lecon.description)
            run_desc.font.name = 'Arial'
            run_desc.font.size = Pt(10.5)
            run_desc.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            run_desc.italic = True

        def formater_run_docx(p, texte_brut):
            # Découpage du texte pour appliquer les styles gras, italique et code
            parties = re.split(r'(\*\*.+?\*\*|\*[^*]+?\*|_[^_]+?_|`[^`]+?`)', texte_brut)
            for partie in parties:
                if not partie:
                    continue
                if partie.startswith('**') and partie.endswith('**'):
                    run = p.add_run(partie[2:-2])
                    run.bold = True
                    run.font.name = 'Arial'
                elif (partie.startswith('*') and partie.endswith('*')) or (partie.startswith('_') and partie.endswith('_')):
                    run = p.add_run(partie[1:-1])
                    run.italic = True
                    run.font.name = 'Arial'
                elif partie.startswith('`') and partie.endswith('`'):
                    run = p.add_run(partie[1:-1])
                    run.font.name = 'Consolas'
                    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                    run.bold = True
                else:
                    run = p.add_run(partie)
                    run.font.name = 'Arial'

        in_blockquote = False
        blockquote_lines = []

        for ligne in contenu_brut.split('\n'):
            ligne_stripped = ligne.strip()

            if in_blockquote and not ligne_stripped.startswith('>'):
                # Création du callout blockquote Word
                table = doc_word.add_table(rows=1, cols=1)
                table.autofit = False
                table.columns[0].width = Cm(16)
                cell = table.cell(0, 0)
                
                # XML shading (fond violet clair) et border (bordure gauche violette)
                quote_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F3FF"/>')
                quote_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="4F46E5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
                cell._tc.get_or_add_tcPr().append(quote_shading)
                cell._tc.get_or_add_tcPr().append(quote_borders)
                
                p_quote = cell.paragraphs[0]
                p_quote.paragraph_format.left_indent = Cm(0.4)
                p_quote.paragraph_format.right_indent = Cm(0.4)
                p_quote.paragraph_format.space_before = Pt(6)
                p_quote.paragraph_format.space_after = Pt(6)
                
                citation_texte = "\n".join(blockquote_lines)
                run_quote = p_quote.add_run(citation_texte)
                run_quote.font.name = 'Arial'
                run_quote.font.size = Pt(10)
                run_quote.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
                run_quote.italic = True
                
                doc_word.add_paragraph()
                blockquote_lines = []
                in_blockquote = False

            if not ligne_stripped:
                if not in_blockquote:
                    doc_word.add_paragraph()
                else:
                    blockquote_lines.append("")
            elif ligne_stripped.startswith('>'):
                in_blockquote = True
                blockquote_lines.append(ligne_stripped.lstrip('>').strip())
            elif ligne_stripped.startswith('### '):
                texte_titre = ligne_stripped[4:].strip()
                p = doc_word.add_heading(level=3)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                run = p.runs[0] if p.runs else p.add_run(texte_titre)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                run.bold = True
            elif ligne_stripped.startswith('## '):
                texte_titre = ligne_stripped[3:].strip()
                p = doc_word.add_heading(level=2)
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
                run = p.runs[0] if p.runs else p.add_run(texte_titre)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                run.bold = True
            elif ligne_stripped.startswith('# '):
                texte_titre = ligne_stripped[2:].strip()
                p = doc_word.add_heading(level=1)
                p.paragraph_format.space_before = Pt(20)
                p.paragraph_format.space_after = Pt(8)
                run = p.runs[0] if p.runs else p.add_run(texte_titre)
                run.font.name = 'Arial'
                run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
                run.bold = True
            elif ligne_stripped.startswith('- ') or ligne_stripped.startswith('* '):
                p = doc_word.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                formater_run_docx(p, ligne_stripped[2:])
            else:
                p = doc_word.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                formater_run_docx(p, ligne_stripped)

        if in_blockquote and blockquote_lines:
            table = doc_word.add_table(rows=1, cols=1)
            table.autofit = False
            table.columns[0].width = Cm(16)
            cell = table.cell(0, 0)
            quote_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F3FF"/>')
            quote_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="4F46E5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
            cell._tc.get_or_add_tcPr().append(quote_shading)
            cell._tc.get_or_add_tcPr().append(quote_borders)
            
            p_quote = cell.paragraphs[0]
            p_quote.paragraph_format.left_indent = Cm(0.4)
            p_quote.paragraph_format.space_before = Pt(6)
            p_quote.paragraph_format.space_after = Pt(6)
            citation_texte = "\n".join(blockquote_lines)
            run_quote = p_quote.add_run(citation_texte)
            run_quote.font.name = 'Arial'
            run_quote.font.size = Pt(10)
            run_quote.font.color.rgb = RGBColor(0x5B, 0x21, 0xB6)
            run_quote.italic = True

        buffer = io.BytesIO()
        doc_word.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{nom_fichier}.docx"'
        return response
    else:
        raise ValueError("Format non supporté. Utilisez 'pdf' ou 'docx'.")
